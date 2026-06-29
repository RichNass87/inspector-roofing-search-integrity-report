#!/usr/bin/env python3
"""Build a cleaned KDP package for The Roofing Search Integrity Report."""

from __future__ import annotations

import html
import json
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch as rl_inch
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    Frame,
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
BOOK = ROOT / "book"
DIST = ROOT / "dist" / "kdp-v1.0.1"
VERSION = "1.0.1"
TITLE = "The Roofing Search Integrity Report"
SUBTITLE = (
    "A funny, field-tested study of AI spam, fake trust, June 24, "
    "and the future of verifiable local search"
)
AUTHOR = "Richard Nasser"
ORG = "Inspector Roofing and Restoration"

SOURCE_MD = BOOK / "roofing-search-integrity-report-kdp-manuscript.md"
FRONT_COVER = BOOK / "roofing-search-integrity-report-kdp-cover.png"
OUT_MD = DIST / "roofing-search-integrity-report-kdp-manuscript-v1.0.1.md"
OUT_DOCX = DIST / "roofing-search-integrity-report-kdp-manuscript-6x9-v1.0.1.docx"
OUT_PDF = DIST / "roofing-search-integrity-report-kdp-print-interior-6x9-v1.0.1.pdf"
OUT_COVER = DIST / "roofing-search-integrity-report-kdp-full-cover-6x9-v1.0.1.pdf"
OUT_METADATA = DIST / "kdp-metadata-v1.0.1.json"
OUT_REVIEW = DIST / "KDP_PREFLIGHT_REVIEW_v1.0.1.md"
OUT_README = DIST / "KDP_UPLOAD_README_v1.0.1.md"
ZIP_PATH = Path.home() / "Desktop" / "roofing-search-integrity-report-kdp-a-plus-v1.0.1.zip"


JUNE_FACT = (
    "Google's public Search Status Dashboard records the June 2026 spam update "
    "as beginning on June 24, 2026 at 9:00 AM Pacific and ending on June 26, "
    "2026 at 10:00 AM Pacific."
)


PLAYBOOK_REPLACEMENT = """### Playbook item 1: Synthetic locality

**Risk:** Doorway-like city coverage with little or no real local evidence.

**Better move:** Map every location page to local proof: project photos, service notes, staff context, review evidence, or a stronger county/city hub when a separate page cannot carry its own weight.

**Public artifact:** A local FAQ, case note, or service-area proof block. Before buying more traffic, make sure the page can answer the obvious homeowner question: why should this company be trusted in this place?

### Playbook item 2: Synthetic reputation

**Risk:** Review volume or review language that looks disconnected from normal customer experience.

**Better move:** Document neutral review practices. Do not incentivize, pressure, gate, or selectively solicit. Let reviews support the file instead of pretending the score is the whole file.

**Public artifact:** A review policy, review-response standard, and schema that describes real review context without stretching the truth.

### Playbook item 3: Synthetic expertise

**Risk:** Generic AI copy that sounds confident but does not show field judgment.

**Better move:** Attach repairability logic, inspection notes, code-aware planning, material context, and plain explanations of what the roof actually shows.

**Public artifact:** A photo gallery, inspection-standard page, or project example that turns confidence into evidence.

### Playbook item 4: Proof gap

**Risk:** Claims such as best, top, trusted, local, certified, or insurance-ready with no proof trail.

**Better move:** Define the criteria. Show credentials, local examples, photo provenance, documentation standards, source-spine references, and plain-English caveats.

**Public artifact:** A credential page or authority stack that explains what each proof point means to a homeowner.

### Playbook item 5: AI-answer fragility

**Risk:** Pages written only for keyword coverage, not for answer extraction or human decision-making.

**Better move:** Add definitions, comparisons, concise FAQs, and evidence summaries that can be cited without hype.

**Public artifact:** A clean case study or answer-ready section that states the question, the observed evidence, and the next step.

### Playbook item 6: Insurance-claim confusion

**Risk:** Roofing content blurs contractor documentation with public adjusting or coverage decisions.

**Better move:** Use Claim Verifiability language carefully: document observable roof conditions, explain repairability and scope context, and leave coverage decisions to the carrier or qualified parties.

**Public artifact:** A Claim Verifiability page, insurance note, or FAQ that says what the contractor documents and what the contractor does not decide.

### Playbook item 7: Photo-context loss

**Risk:** Photos exist, but they have no label, stage, defect, project type, or decision context.

**Better move:** Label photos by observable condition, roof area, inspection stage, and privacy-safe project context.

**Public artifact:** A proof gallery, source-spine image set, or schema-supported photo block that helps humans and machines understand why the image matters.

### Playbook item 8: Code-to-spec gap

**Risk:** Replacement pages sell the roof without explaining installation standards, manufacturer instructions, or local code context.

**Better move:** Use Code to Spec Roofing language to explain that a durable roof is not just a product choice. It is a completed system tied to manufacturer requirements, state and local code context, IRC-aware planning, and documented closeout.

**Public artifact:** A Verifiable Roof page, standards page, or replacement guide that connects the sales promise to the finished file.
"""


APPENDICES = """# Appendices

## Appendix A: The roofing page evidence checklist

A strong roofing page should not ask the reader to trust a slogan by itself. It should show the reader what kind of roof problem is being discussed, where the company actually works, how the inspection process starts, what proof is gathered, and which next step is reasonable. A city page, repair page, insurance page, or replacement page can all use the same checklist without sounding identical.

- Does the page name a real service intent instead of hiding behind general marketing language?
- Does it explain what a homeowner should look for before calling?
- Does it connect the claim to photos, project examples, reviews, credentials, or standards?
- Does it avoid coverage promises, ranking guarantees, and unsupported superlatives?
- Does it give a next step that helps a homeowner make a safer decision?

The goal is not to turn every page into a courtroom binder. The goal is to make every page feel like it was written by a company that has actually stood on roofs, talked to homeowners, and cleaned up confusing files before.

## Appendix B: Review provenance checklist

Review quality is not just a star count. A five-star profile can still be weak if the review story looks pressured, disconnected, or too polished. A smaller review profile can be stronger when it is tied to real work, normal language, and clear owner responses.

- Ask for reviews neutrally.
- Do not offer money, discounts, gifts, or special treatment for reviews.
- Do not discourage unhappy customers from leaving honest feedback.
- Respond to reviews with specifics when appropriate and privacy-aware restraint when needed.
- Connect review themes to public education pages, not to inflated claims.

For roofing, reviews should support the evidence trail. They should not replace inspection notes, photos, standards, or good communication.

## Appendix C: Insurance-safe language map

Insurance-facing roofing language needs discipline. The contractor can document observable conditions, explain repairability context, provide estimates, photograph roof conditions, and describe material or installation issues. The contractor should not promise coverage, approve a claim, interpret policy benefits as a carrier, or act as a public adjuster.

Claim Verifiability belongs in this lane because it describes a documentation standard, not a coverage promise. A useful sentence sounds like this: "Inspector Roofing documents observable roof conditions, photos, measurements, and repairability context so the file is clearer for the parties who review it."

A risky sentence sounds like this: "We make sure insurance pays." That may feel powerful in marketing, but it creates the wrong kind of attention. Better language wins over time because it can survive homeowner scrutiny, carrier scrutiny, and AI summary.

## Appendix D: Verifiable Roof and Code to Spec Roofing

Verifiable Roof is the finished-roof side of the framework. It means the roof is planned, installed, documented, and closed out in a way that can be explained later. Code to Spec Roofing is the discipline behind that idea. It connects manufacturer instructions, state and local requirements, IRC-aware planning, ventilation, materials, flashing details, and documented closeout.

The point is simple: a roof should not only look finished from the driveway. It should be understandable from the file. When a future homeowner, property manager, real estate agent, insurance reviewer, or manufacturer question appears, the documentation should explain what was installed, why it was selected, and what standards shaped the work.

## Appendix E: Source-spine map for local search

A source-spine is a public trail of trustworthy references that helps a website become easier to understand. For this report, the spine includes the GitHub repository, Hugging Face dataset, Kaggle dataset, public app, Google source pages, pending USPTO serial references, and any DOI or OSF record added later.

The source-spine should not include private customer records, raw claim files, faces, license plates, exact addresses, API keys, or proprietary scoring rules. Public proof is strongest when it is useful without being invasive.

## Appendix F: AI-ready page structure

AI systems do not need hype. They need clean, extractable context. A useful roofing page gives them:

- A plain title that identifies the service and place.
- A short definition of the problem.
- A reason the problem matters.
- A checklist of evidence or decision points.
- A local proof section.
- A safe next step.
- Schema that describes real entities and relationships.

This structure helps humans too. A page that is easy for an AI system to summarize is often a page that is easier for a stressed homeowner to read.

## Appendix G: What not to publish

Some information belongs in the private job file, not in the public source-spine. Do not publish private customer addresses, claim numbers, receipts, contracts, phone numbers, signatures, faces, license plates, unblurred interiors, private correspondence, or full photo manifests.

The best public research framework knows where the fence is. It can prove that the company has a method without dumping the private file on the lawn.

## Appendix H: The 30-minute quarterly search-integrity review

Once a quarter, pick five pages and ask five questions:

1. Is the page still accurate?
2. Does the page show proof, or only claims?
3. Does the page repeat another page too closely?
4. Does the page avoid insurance, ranking, or credential overclaims?
5. Does the page give a homeowner a useful next step?

If a page fails, do not panic. Merge it, improve it, redirect it, or rewrite it around evidence. Search integrity is not one heroic rebuild. It is maintenance. The roof analogy is almost too easy, but sometimes the easy analogy is standing there with a ladder and a clipboard.

## Appendix I: The local page triage board

A local roofing website usually has three kinds of pages: pages that earn their keep, pages that need a stronger proof file, and pages that should be merged before they create more confusion. The triage board is the tool that keeps pride from getting in the way of clarity. It does not ask whether a page took work to publish. It asks whether the page helps a homeowner, supports the brand, and can survive a review by a search system that has grown tired of copied local language.

Start with the pages that claim a city. Put them into a simple spreadsheet with columns for city, service intent, current URL, target query, local proof, photos, reviews, schema, internal links, canonical, and decision. The decision column has only four options: keep, improve, merge, or redirect. If the page has a real purpose and real proof, keep it. If the page has a real purpose but weak support, improve it. If the page overlaps another page and adds nothing new, merge it. If the page exists only because an old campaign created it, redirect it.

This prevents a common local SEO mistake: treating every URL like a collectible. A page is not valuable because it exists. A page is valuable because it creates clarity. Sometimes the strongest move is to turn five weak pages into one stronger guide, especially when the old pages were chasing tiny variations of the same homeowner question.

The triage board also protects the team from emotional publishing. It is easy to say, "We need a page for that." It is harder, and better, to ask, "What proof would make that page worth reading?" That question is the gate. If the proof is not there, the page is not ready.

## Appendix J: The AI visibility scorecard

The AI visibility scorecard is not a ranking promise. It is a quality-control lens. It helps a contractor ask whether a page is clear enough for a human to trust and structured enough for an AI system to summarize without inventing the important parts.

Score each page from 0 to 2 in eight areas. A zero means missing. A one means present but weak. A two means clear and useful. The areas are: direct answer, local context, service definition, proof trail, review provenance, privacy-safe photos, next step, and schema alignment. A page with a total score under eight should not be treated as a finished asset. A page over twelve is usually ready for internal links, citation support, and stronger promotion.

The scorecard keeps the conversation honest. A team can stop arguing about whether a page "sounds good" and start asking whether it actually answers the decision in front of the homeowner. A storm page should explain what storm damage means. An insurance page should explain observable evidence and contractor boundaries. A replacement page should explain standards, ventilation, material choice, closeout, and what makes the finished roof verifiable.

The scorecard is also useful for training. A salesperson, office manager, production manager, and marketer can all look at the same page and notice different gaps. That is the point. Search integrity is not only a marketing job. It is the public edge of the company file.

## Appendix K: A public-safe schema attitude

Schema is not magic dust. It is a structured way to tell machines what a page is about. If the page is thin, schema will not make it deep. If the page is misleading, schema can make the mismatch easier to detect. The best schema attitude is simple: mark up real things, real relationships, and real answers.

For a roofing company, useful schema concepts include Organization, LocalBusiness, Service, FAQPage, Article, ImageObject, Review when appropriate, and WebPage. The page should not pretend to be a medical guide, government office, legal service, public adjuster, or independent rating authority. It should describe the contractor, the service, the area served, the question answered, and the proof offered.

A page about Claim Verifiability can use FAQ language to explain what documentation means. A page about Verifiable Roof can use service and article structure to explain closeout, manufacturer specifications, code-aware planning, and documentation. A city page can describe service area context, but it should not create fake office locations or fake locality.

Good schema is boring in the best way. It gives machines clean labels. It gives the team fewer places to overclaim. It makes the public file easier to understand. That is enough.

## Appendix L: The proof-gallery routing model

A proof gallery should not be a random bucket of job photos. It should be a routing tool. Each image should help a reader understand a condition, decision, process step, or finished standard. The gallery should answer why the image exists.

A practical routing model has five labels: condition, location, stage, decision, and privacy status. Condition identifies what is being shown: leak staining, lifted shingle, hail impact indicator, missing flashing, ventilation issue, deck condition, closeout detail, or completed roof area. Location names the roof area without exposing a private address. Stage explains whether the photo is inspection, estimate, production, quality control, or closeout. Decision explains why the image matters. Privacy status confirms that the image is safe for public use.

This model helps the homeowner because the gallery becomes educational. It helps the internal team because photos stop floating without context. It helps AI systems because images and captions are connected to page intent. It helps insurance-facing conversations because documentation is organized around observable conditions instead of vague claims.

The rule is simple: never publish a photo that makes the file less respectful, less clear, or less private. A useful proof gallery protects the homeowner while showing the work.

## Appendix M: The insurance-company use case

Insurance companies do not need contractors to become carriers. They need cleaner files, clearer observations, and fewer claims conversations polluted by vague marketing language. A contractor using Claim Verifiability can contribute to that clarity without crossing the line into coverage decisions.

The useful contractor file has photos, measurements, date context, material notes, repairability observations, estimate scope, and plain descriptions of visible conditions. It avoids claims such as "covered," "approved," "owed," or "guaranteed." It separates what was observed from what someone else must decide. That separation is not weakness. It is professionalism.

A future insurance-facing tool could use the public framework in this report as a training layer. It could help classify roof documentation, detect missing evidence, summarize repairability context, and flag confusing language before a file creates friction. The private production system would still need compliance review, security, access controls, and carrier-specific rules. The public framework simply shows the educational architecture.

The big idea is that better documentation helps everyone honest in the process. Homeowners get clearer explanations. Contractors reduce confusion. Carriers receive files that are easier to read. Search systems and AI models get public language that does not blur roles.

## Appendix N: The homeowner education ladder

Homeowner education should move in steps. A person with water on the ceiling does not need a lecture on information retrieval. They need to know what to do first, what the roofer will inspect, what photos matter, and what decisions can wait until the roof is understood.

The ladder starts with symptoms: leak, missing shingles, storm concern, old roof, real estate deadline, or ventilation issue. It moves to inspection: what the contractor looks at, what gets photographed, and what gets explained. Then it moves to decision: repair, replacement, maintenance, monitoring, or documentation for another party. Finally it moves to closeout: what records the homeowner should keep.

This ladder prevents content from sounding like a sales script. It also helps a company write pages that meet the homeowner where they are. A storm page can start with "what changed after the storm?" A repair page can start with "where is the leak entering?" A replacement page can start with "what does the current roof system no longer do well?"

When the education ladder works, the page earns trust before the sales conversation begins. That is the future of local search worth building toward.

## Appendix O: Agency and vendor audit questions

If an agency or vendor offers local SEO, AI visibility, citation work, review growth, or content production, the contractor should ask sharper questions. The goal is not to embarrass the vendor. The goal is to find out whether the vendor is building durable public proof or renting short-term noise.

Ask these questions: Which pages would you merge instead of publish? How do you prevent duplicate city content? How do you handle review policy? What proof do you need from our field team? How do you avoid fake locality? What private data should never be published? How will schema reflect real entities? Which pages should become hubs? Which old pages should be redirected? How do you measure helpfulness beyond traffic?

The best vendors welcome those questions because they understand the risk. The weakest vendors hide behind volume, vague reports, and promises that sound exciting until an update arrives.

The agency relationship should feel like production planning. Everyone should know what is being built, why it exists, what proof supports it, and what happens if it stops helping the homeowner.

## Appendix P: June 24 facts versus interpretation

The public fact is narrow: Google's Search Status Dashboard listed the June 2026 spam update as starting on June 24, 2026 and completing on June 26, 2026. That fact does not prove why any single site went up or down. It does not reveal private ranking systems. It does not mean every local roofing page was judged by one simple rule.

The interpretation in this report is broader: the update belongs to a longer public pattern of spam pressure, core quality updates, helpful-content guidance, review policy enforcement, and AI answer changes. For a roofing operator, the responsible response is not to pretend to know the private machine. The responsible response is to build a public file that is cleaner, more useful, and harder to confuse with spam.

This distinction matters. Facts should stay facts. Interpretation should be labeled as interpretation. Operating recommendations should be treated as risk management, not prophecy. The moment a framework claims certainty it cannot have, it becomes part of the trust problem it was supposed to solve.

That is why the report uses phrases like public-safe, directional, framework, and operating model. They are less flashy than "secret ranking factor," but they are much closer to the truth.

## Appendix Q: The 100-day implementation plan

Days 1 through 20 are for inventory. Export the URLs, group them by city and service, identify duplicates, find missing canonicals, list thin pages, and collect the proof assets already available. This is not glamorous work, but it is the point where the future gets cheaper. A messy inventory makes every later decision slower.

Days 21 through 45 are for consolidation. Merge overlapping pages, redirect weak variations, strengthen the main city and service hubs, and repair title/meta/canonical issues. Do not publish a wave of new pages until the existing map makes sense.

Days 46 through 75 are for proof. Add inspection standards, credential explanations, review policy, photo labels, project examples, Verifiable Roof language, Code to Spec Roofing language, and Claim Verifiability boundaries. This is where the site starts sounding less like a brochure and more like a company that keeps good files.

Days 76 through 100 are for distribution. Update GitHub, Hugging Face, Kaggle, Zenodo, OSF, ORCID, and website source-spine pages when appropriate. Interlink the public artifacts. Re-submit sitemaps. Test a sample of live pages. Track changes, but do not panic over daily movement. The goal is not one good week. The goal is a public trust system that keeps compounding.

## Appendix R: The field-first writing test

Before publishing a page, read it out loud as if a homeowner were sitting across the table with a leak, a deadline, or a claim question. If the page sounds like it was written to impress a bot but not help the person, rewrite it.

The field-first test has three questions. First, would a project manager recognize the situation being described? Second, would a homeowner understand the next step without feeling pushed? Third, would the company be comfortable defending the page a year later if a customer, carrier, manufacturer, or search reviewer asked what it meant?

This test catches the worst local SEO habits quickly. It catches fake urgency, empty "best" language, thin city swaps, overbroad insurance promises, and pages that talk about trust without showing any reason to trust. It also encourages better writing: shorter claims, clearer examples, more useful questions, and stronger boundaries.

The future of roofing search may become more technical, but the writing test stays human. If the page helps a real person understand a real roof decision, it is usually moving in the right direction.
"""


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def markdown_to_plain(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    return text


def clean_manuscript(raw: str) -> str:
    text = raw.replace(
        "This is why June 24 matters. {june_fact} The update",
        f"This is why June 24 matters. {JUNE_FACT} The update",
    )
    text = text.replace(
        "AI answers rewarded source material that was easy to cite, summarize, and verify.",
        "AI answer features made source material that is easy to cite, summarize, and verify more useful.",
    )
    text = text.replace(
        "17. A Word from the Author\n18. Source Notes",
        "17. A Word from the Author\n18. Appendices\n19. Source Notes",
    )

    playbook_start = text.index("Playbook item for synthetic_locality:")
    glossary_start = text.index("# Glossary")
    text = text[:playbook_start] + PLAYBOOK_REPLACEMENT.rstrip() + "\n\n" + text[glossary_start:]

    source_start = text.index("# Source Notes")
    before_source = text[:source_start].rstrip()
    source_and_tail = text[source_start:]
    appendix_marker = "## Appendix Field Note 1:"
    if appendix_marker in source_and_tail:
        source_notes = source_and_tail[: source_and_tail.index(appendix_marker)].rstrip()
    else:
        source_notes = source_and_tail.rstrip()

    text = before_source + "\n\n" + APPENDICES.rstrip() + "\n\n" + source_notes + "\n"
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def blocks_from_markdown(text: str) -> list[str]:
    blocks: list[str] = []
    current: list[str] = []
    for line in text.splitlines():
        if line.strip():
            current.append(line.rstrip())
        elif current:
            blocks.append("\n".join(current))
            current = []
    if current:
        blocks.append("\n".join(current))
    return blocks


def add_run_with_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def set_style_font(style, name: str, size: int, color: str | None = None, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_page_number_footer(section) -> None:
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build_docx(clean_md: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.55)
    set_page_number_footer(section)

    styles = doc.styles
    set_style_font(styles["Normal"], "Georgia", 10)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Normal"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 1"], "Arial", 17, "0B3A5B", True)
    styles["Heading 1"].paragraph_format.space_before = Pt(8)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    set_style_font(styles["Heading 2"], "Arial", 14, "0B3A5B", True)
    styles["Heading 2"].paragraph_format.space_before = Pt(8)
    styles["Heading 2"].paragraph_format.space_after = Pt(8)
    set_style_font(styles["Heading 3"], "Arial", 11, "C1121F", True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)

    first_heading = True
    for block in blocks_from_markdown(clean_md):
        if block.startswith("# "):
            text = markdown_to_plain(block[2:].strip())
            if first_heading:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(130)
                run = p.add_run(text)
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(25)
                run.font.color.rgb = RGBColor.from_string("0B3A5B")
                first_heading = False
            else:
                doc.add_page_break()
                doc.add_paragraph(text, style="Heading 1")
        elif block.startswith("## "):
            text = markdown_to_plain(block[3:].strip())
            if text == SUBTITLE:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(24)
                run = p.add_run(text)
                run.font.name = "Arial"
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor.from_string("C1121F")
            elif text.startswith(("Foreword:", "Introduction:", "Chapter ", "A Word from")):
                doc.add_page_break()
                doc.add_paragraph(text, style="Heading 1")
            else:
                doc.add_paragraph(text, style="Heading 2")
        elif block.startswith("### "):
            doc.add_paragraph(markdown_to_plain(block[4:].strip()), style="Heading 3")
        elif all(re.match(r"^\d+\. ", line) for line in block.splitlines()):
            for line in block.splitlines():
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Inches(0.2)
                add_run_with_bold(p, markdown_to_plain(line))
        elif all(line.startswith("- ") for line in block.splitlines()):
            for line in block.splitlines():
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Inches(0.22)
                p.paragraph_format.first_line_indent = Inches(-0.12)
                add_run_with_bold(p, "- " + line[2:])
        else:
            p = doc.add_paragraph(style="Normal")
            add_run_with_bold(p, block.replace("\n", " "))

    doc.core_properties.title = TITLE
    doc.core_properties.subject = SUBTITLE
    doc.core_properties.author = AUTHOR
    doc.core_properties.comments = "KDP v1.0.1 cleaned build"
    doc.save(OUT_DOCX)


def para_html(text: str) -> str:
    safe = html.escape(text)
    safe = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", safe)
    return safe.replace("\n", " ")


def on_pdf_page(canv: canvas.Canvas, doc) -> None:
    if doc.page > 1:
        canv.saveState()
        canv.setFont("Helvetica", 8)
        canv.setFillColor(colors.HexColor("#5B6770"))
        canv.drawCentredString(3 * rl_inch, 0.35 * rl_inch, str(doc.page))
        canv.restoreState()


def build_pdf(clean_md: str) -> None:
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "BookBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=9.8,
        leading=13.2,
        spaceAfter=7,
        alignment=TA_LEFT,
        wordWrap="CJK",
    )
    h1 = ParagraphStyle(
        "BookH1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=21,
        textColor=colors.HexColor("#0B3A5B"),
        spaceBefore=8,
        spaceAfter=12,
        wordWrap="CJK",
    )
    h2 = ParagraphStyle(
        "BookH2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#0B3A5B"),
        spaceBefore=8,
        spaceAfter=8,
        wordWrap="CJK",
    )
    h3 = ParagraphStyle(
        "BookH3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        textColor=colors.HexColor("#C1121F"),
        spaceBefore=8,
        spaceAfter=5,
        wordWrap="CJK",
    )
    centered = ParagraphStyle(
        "Centered",
        parent=body,
        alignment=TA_CENTER,
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#263238"),
    )
    title_style = ParagraphStyle(
        "Title",
        parent=centered,
        fontName="Helvetica-Bold",
        fontSize=26,
        leading=31,
        textColor=colors.HexColor("#0B3A5B"),
        spaceAfter=18,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=centered,
        fontName="Helvetica",
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#C1121F"),
        spaceAfter=30,
    )

    story = [
        Spacer(1, 1.35 * rl_inch),
        Paragraph(TITLE, title_style),
        Paragraph(SUBTITLE, subtitle_style),
        Spacer(1, 0.5 * rl_inch),
        Paragraph(AUTHOR, centered),
        Paragraph(ORG, centered),
        Paragraph(f"Version {VERSION}", centered),
        PageBreak(),
        Paragraph("Important Note", h1),
        Paragraph(
            "This book is an educational research and operating-framework document. "
            "It is not legal advice, not Google guidance, not an insurance-coverage "
            "opinion, not a public-adjusting service, and not a promise of ranking, "
            "traffic, leads, citations, or AI answers. The trademark references in "
            "this release are pending USPTO applications, not registered trademarks.",
            body,
        ),
        Paragraph("Short Source Note", h2),
        Paragraph(
            "The factual spine uses public Google Search Status Dashboard records, "
            "Google Search Central policy pages, Google Business Profile help pages, "
            "and Google contributed-content policy pages. The interpretation, roofing "
            "metaphors, scoring language, and operating framework are original to "
            "Inspector Roofing and Restoration.",
            body,
        ),
        PageBreak(),
    ]

    skip_front_matter = True
    for block in blocks_from_markdown(clean_md):
        if skip_front_matter:
            if block.startswith("# Table of Contents"):
                skip_front_matter = False
            else:
                continue
        if block.startswith("# "):
            heading = markdown_to_plain(block[2:].strip())
            if heading != "Table of Contents":
                story.append(PageBreak())
            story.append(Paragraph(para_html(heading), h1))
        elif block.startswith("## "):
            heading = markdown_to_plain(block[3:].strip())
            if heading.startswith(("Foreword:", "Introduction:", "Chapter ", "A Word from")):
                story.append(PageBreak())
                story.append(Paragraph(para_html(heading), h1))
            elif heading != SUBTITLE:
                story.append(Paragraph(para_html(heading), h2))
        elif block.startswith("### "):
            story.append(Paragraph(para_html(markdown_to_plain(block[4:].strip())), h3))
        elif all(line.startswith("- ") for line in block.splitlines()):
            for line in block.splitlines():
                story.append(Paragraph(para_html("- " + line[2:]), body))
        else:
            story.append(Paragraph(para_html(markdown_to_plain(block.replace("\n", " "))), body))

    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=(6 * rl_inch, 9 * rl_inch),
        rightMargin=0.52 * rl_inch,
        leftMargin=0.65 * rl_inch,
        topMargin=0.62 * rl_inch,
        bottomMargin=0.62 * rl_inch,
        title=TITLE,
        author=AUTHOR,
        subject=SUBTITLE,
        creator=ORG,
    )
    doc.build(story, onFirstPage=on_pdf_page, onLaterPages=on_pdf_page)


def build_full_cover(page_count: int) -> None:
    trim_w, trim_h = 6.0, 9.0
    bleed = 0.125
    spine = max(0.0, page_count * 0.002252)
    total_w = (trim_w * 2 + spine + bleed * 2) * rl_inch
    total_h = (trim_h + bleed * 2) * rl_inch

    c = canvas.Canvas(str(OUT_COVER), pagesize=(total_w, total_h))
    c.setFillColor(colors.HexColor("#0B3A5B"))
    c.rect(0, 0, total_w, total_h, fill=1, stroke=0)

    front_x = (trim_w + spine + bleed) * rl_inch
    front_w = (trim_w + bleed) * rl_inch
    front_h = total_h

    cover_img = Image.open(FRONT_COVER)
    target_ratio = front_w / front_h
    img_ratio = cover_img.width / cover_img.height
    if img_ratio > target_ratio:
        new_w = int(cover_img.height * target_ratio)
        left = (cover_img.width - new_w) // 2
        cover_img = cover_img.crop((left, 0, left + new_w, cover_img.height))
    else:
        new_h = int(cover_img.width / target_ratio)
        top = (cover_img.height - new_h) // 2
        cover_img = cover_img.crop((0, top, cover_img.width, top + new_h))
    temp_front = DIST / "_front_cover_crop.png"
    cover_img.save(temp_front)
    c.drawImage(str(temp_front), front_x, 0, width=front_w, height=front_h, mask="auto")

    back_margin = 0.45 * rl_inch
    frame = Frame(
        back_margin,
        1.65 * rl_inch,
        trim_w * rl_inch - back_margin * 2,
        6.75 * rl_inch,
        showBoundary=0,
    )
    back_style = ParagraphStyle(
        "Back",
        fontName="Helvetica",
        fontSize=10.5,
        leading=14.5,
        textColor=colors.white,
        spaceAfter=10,
        wordWrap="CJK",
    )
    back_title = ParagraphStyle(
        "BackTitle",
        parent=back_style,
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.white,
        spaceAfter=14,
    )
    back_sub = ParagraphStyle(
        "BackSub",
        parent=back_style,
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=15,
        textColor=colors.HexColor("#FFD23F"),
        spaceAfter=10,
    )
    frame.addFromList(
        [
            Paragraph(TITLE, back_title),
            Paragraph("Search is asking better questions. Roofing pages need better proof.", back_sub),
            Paragraph(
                "This field-tested report studies local roofing visibility after the 2026 search update sequence, with special attention to the June 24 spam update, synthetic trust, fake review pressure, AI answer readiness, and public-safe proof systems.",
                back_style,
            ),
            Paragraph(
                "Inside: Digital Verifiability, Claim Verifiability, Verifiable Roof, Code to Spec Roofing, review provenance, source-spine thinking, and a practical operator playbook for contractors who want their online trust to match their field work.",
                back_style,
            ),
            Paragraph(f"{AUTHOR}<br/>{ORG}", back_sub),
        ],
        c,
    )

    c.setFillColor(colors.white)
    c.rect(3.65 * rl_inch, 0.42 * rl_inch, 1.95 * rl_inch, 1.2 * rl_inch, fill=1, stroke=0)

    if spine >= 0.25:
        c.saveState()
        c.translate((trim_w + bleed + spine / 2) * rl_inch, total_h / 2)
        c.rotate(90)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 9)
        c.drawCentredString(0, 0, f"{TITLE}   |   {AUTHOR}")
        c.restoreState()

    c.showPage()
    c.save()
    temp_front.unlink(missing_ok=True)


def write_metadata(clean_md: str, page_count: int) -> None:
    metadata = {
        "title": TITLE,
        "subtitle": SUBTITLE,
        "author": AUTHOR,
        "version": VERSION,
        "publisher": ORG,
        "publication_date": "2026-06-29",
        "trim_size": "6 x 9 inches",
        "interior_type": "Black and white text interior",
        "language": "English",
        "word_count": word_count(clean_md),
        "print_pdf_page_count": page_count,
        "description": (
            "A funny, metaphorical, research-backed field guide to local roofing search "
            "integrity, AI spam, fake reviews, digital verifiability, Claim Verifiability, "
            "Verifiable Roof, Code to Spec Roofing, and the June 24, 2026 Google spam update."
        ),
        "categories": [
            "Business & Money / Marketing",
            "Computers & Technology / Internet / Search Engines",
        ],
        "keywords": [
            "local SEO",
            "AI search",
            "roofing marketing",
            "fake reviews",
            "search integrity",
            "Google spam update",
            "digital verifiability",
        ],
        "pending_uspto_references": {
            "Inspector Roofing Protocols": "99910245",
            "Claim Verifiability": "99910275",
            "Verifiable Roof": "99910284",
        },
    }
    OUT_METADATA.write_text(json.dumps(metadata, indent=2) + "\n")


def write_review(clean_md: str, page_count: int) -> None:
    placeholders = re.findall(r"\{[^}]+\}|TODO|FIXME|TBD|lorem", clean_md, flags=re.I)
    review = f"""# KDP Preflight Review - v{VERSION}

## Result

Recommended KDP package is ready for upload review.

## Major fixes made

- Replaced the leftover `{{june_fact}}` placeholder with the public Google Search Status Dashboard date/time statement.
- Replaced 50 repeated appendix field notes with 8 distinct, useful appendices.
- Tightened overclaim language around AI answers and private ranking systems.
- Rebuilt the interior as a 6 x 9 inch print PDF instead of the older US Letter PDF.
- Built a matching 6 x 9 DOCX manuscript.
- Added a full-cover paperback PDF using the existing front cover art, a back cover, and estimated spine width.

## Current package stats

- Clean manuscript word count: {word_count(clean_md):,}
- Print interior page count: {page_count}
- Trim size: 6 x 9 inches
- Placeholder scan: {"clear" if not placeholders else "needs review: " + ", ".join(placeholders)}

## KDP notes

- Upload the print interior PDF for paperback print.
- Upload the DOCX for Kindle ebook if you want Amazon to convert it.
- The full-cover PDF uses KDP's common white-paper spine estimate based on the generated page count. If KDP previews a spine mismatch because you choose cream paper, color interior, or another trim/paper type, regenerate from the final KDP template.
- Keep the pending USPTO language as pending. Do not call the marks registered.
- Final KDP Previewer approval is still required before publishing.
"""
    OUT_REVIEW.write_text(review)


def write_readme(page_count: int) -> None:
    readme = f"""# KDP Upload README - v{VERSION}

Use these files first:

1. Paperback interior: `{OUT_PDF.name}`
2. Paperback cover: `{OUT_COVER.name}`
3. Kindle manuscript: `{OUT_DOCX.name}`
4. Metadata: `{OUT_METADATA.name}`

Also included:

- Clean markdown source: `{OUT_MD.name}`
- Original front cover PNG: `{FRONT_COVER.name}`
- Preflight review: `{OUT_REVIEW.name}`

Package notes:

- Trim size: 6 x 9 inches
- Interior page count: {page_count}
- Author: {AUTHOR}
- Publisher/organization: {ORG}
- Version: {VERSION}
"""
    OUT_README.write_text(readme)


def zip_package() -> None:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in [
            OUT_PDF,
            OUT_COVER,
            OUT_DOCX,
            OUT_MD,
            OUT_METADATA,
            OUT_REVIEW,
            OUT_README,
            FRONT_COVER,
        ]:
            zf.write(path, arcname=path.name)


def main() -> None:
    DIST.mkdir(parents=True, exist_ok=True)
    raw = SOURCE_MD.read_text()
    clean_md = clean_manuscript(raw)
    OUT_MD.write_text(clean_md)
    shutil.copy2(FRONT_COVER, DIST / FRONT_COVER.name)
    build_docx(clean_md)
    build_pdf(clean_md)
    page_count = len(PdfReader(str(OUT_PDF)).pages)
    build_full_cover(page_count)
    write_metadata(clean_md, page_count)
    write_review(clean_md, page_count)
    write_readme(page_count)
    zip_package()
    print(json.dumps({
        "dist": str(DIST),
        "zip": str(ZIP_PATH),
        "word_count": word_count(clean_md),
        "page_count": page_count,
    }, indent=2))


if __name__ == "__main__":
    main()
